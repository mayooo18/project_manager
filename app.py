from functools import wraps
import time
from flask import Flask, render_template, redirect, request, url_for, flash, session
from flask_wtf.csrf import CSRFProtect, CSRFError
from extensions import db, limiter
from forms import WorkerForm, ProjectForm, FileUploadForm, WorkLogForm, WorkLogFilterForm, PaymentForm, PaymentFilterForm, ExpenseForm, IncomeForm, LoginForm
from models import  Project, ProjectFile, WorkLog, Worker, Payment, Expense, Income, Customer, Location, find_or_create_location
from werkzeug.utils import secure_filename
import os
import uuid
from flask_login import LoginManager, login_required, current_user, login_user, logout_user
import logging 
from datetime import timedelta, datetime
from sqlalchemy import text
from sqlalchemy.exc import OperationalError, DisconnectionError
from document_forms import DocumentForm
from ai_routes import ai_bp
from api_routes import api_bp as field_api_bp
from google_routes import google_bp
from reminder_routes import reminder_bp
from vehicle_routes import vehicle_bp
from quote_routes import quote_bp
from invoice_routes import invoice_bp
from contract_routes import contract_bp
from change_order_routes import change_order_bp
from subcontractor_routes import subcontractor_bp
from license_routes import license_bp
from permit_routes import permit_bp
from waiver_routes import waiver_bp
from portal_routes import portal_bp
from crew_routes import crew_bp
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from flask import send_file  



app = Flask(__name__)
app.config.from_object("config.Config")

db.init_app(app)
csrf = CSRFProtect(app)
limiter.init_app(app)
app.register_blueprint(ai_bp)
app.register_blueprint(field_api_bp)
app.register_blueprint(google_bp)
app.register_blueprint(reminder_bp)
app.register_blueprint(vehicle_bp)
app.register_blueprint(quote_bp)
app.register_blueprint(invoice_bp)
app.register_blueprint(contract_bp)
app.register_blueprint(change_order_bp)
app.register_blueprint(subcontractor_bp)
app.register_blueprint(license_bp)
app.register_blueprint(permit_bp)
app.register_blueprint(waiver_bp)
app.register_blueprint(portal_bp)
app.register_blueprint(crew_bp)

# API routes use API key auth — exempt from CSRF
csrf.exempt(field_api_bp)
csrf.exempt(google_bp)

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_UPLOAD_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'gif', 'heic', 'doc', 'docx', 'xls', 'xlsx', 'csv', 'txt'}


def save_uploaded_file(file_storage):
    """Validate extension and save with a randomized filename. Returns the
    stored filename, or None if the file type is not allowed."""
    original = secure_filename(file_storage.filename)
    ext = original.rsplit('.', 1)[-1].lower() if '.' in original else ''
    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        return None
    filename = f"{uuid.uuid4().hex}_{original}"
    file_storage.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    return filename


@app.after_request
def set_security_headers(response):
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['Permissions-Policy'] = (
        'camera=(), microphone=(), geolocation=(), payment=()')
    response.headers['Content-Security-Policy'] = '; '.join([
        "default-src 'self'",
        "script-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com",
        "style-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://fonts.googleapis.com",
        "font-src 'self' data: https://fonts.gstatic.com",
        "img-src 'self' data: blob:",
        "connect-src 'self'",
        "object-src 'none'",
        "base-uri 'self'",
        "form-action 'self'",
        "frame-ancestors 'none'",
    ])
    if app.config.get('ENVIRONMENT') == 'production':
        response.headers['Strict-Transport-Security'] = (
            'max-age=31536000; includeSubDomains')
    return response


from models import Worker , User


login_manager = LoginManager()
login_manager.login_view = 'login'  # Redirect to login page if not authenticated
login_manager.init_app(app)



@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


from permissions import owner_required


def _owner_count():
    return User.query.filter_by(role='owner', active=True).count()


@app.route('/users', methods=['GET', 'POST'])
@login_required
@owner_required
def users():
    from forms import UserAdminForm
    form = UserAdminForm()
    if form.validate_on_submit():
        if User.query.filter_by(username=form.username.data).first():
            flash('That username is already taken.', 'error')
        else:
            u = User(username=form.username.data, role='admin', active=True)
            u.set_password(form.password.data)
            db.session.add(u)
            db.session.commit()
            flash('Admin (foreman) account created.')
        return redirect(url_for('users'))
    all_users = User.query.order_by(User.role.desc(), User.username.asc()).all()
    return render_template('users.html', form=form, users=all_users)


@app.route('/users/<int:user_id>/role', methods=['POST'])
@login_required
@owner_required
def user_role(user_id):
    u = User.query.get_or_404(user_id)
    new_role = request.form.get('role')
    if new_role not in ('owner', 'admin'):
        flash('Invalid role.', 'error')
    elif u.role == 'owner' and new_role == 'admin' and _owner_count() <= 1:
        flash('You cannot demote the last owner.', 'error')
    else:
        u.role = new_role
        db.session.commit()
        flash(f'{u.username} is now {new_role}.')
    return redirect(url_for('users'))


@app.route('/users/<int:user_id>/toggle-active', methods=['POST'])
@login_required
@owner_required
def user_toggle_active(user_id):
    u = User.query.get_or_404(user_id)
    if u.active and u.role == 'owner' and _owner_count() <= 1:
        flash('You cannot disable the last owner.', 'error')
    else:
        u.active = not u.active
        db.session.commit()
        flash(f"{u.username} {'enabled' if u.active else 'disabled'}.")
    return redirect(url_for('users'))

@app.route('/login', methods=['GET', 'POST'])
@limiter.limit("5 per minute")
def login():
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and not user.active:
            flash('That account has been disabled. Ask the owner to re-enable it.')
            return redirect(url_for('login'))
        if user and user.check_password(form.password.data):
            login_user(user)
            session.permanent = True
            flash('Logged in successfully.')
            return redirect(url_for('home'))
        else:
            flash('Invalid username or password.')
            return redirect(url_for('login'))
            
    return render_template('login.html', form=form)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Logged out successfully.')
    return redirect(url_for('login'))


@app.route('/')
@login_required
def home():
    from sqlalchemy import func
    from datetime import date
    from models import Customer, Quote, Reminder

    worker_count = Worker.query.count()
    project_count = Project.query.count()

    # ── Dashboard metrics ──
    active_projects = Project.query.filter_by(status='Active').count()
    active_workers = Worker.query.filter_by(active=True).count()
    customer_count = Customer.query.filter_by(user_id=current_user.id).count()

    total_income = db.session.query(func.coalesce(func.sum(Income.amount), 0.0)).scalar() or 0.0
    total_expenses = db.session.query(func.coalesce(func.sum(Expense.amount), 0.0)).scalar() or 0.0
    total_payments = db.session.query(func.coalesce(func.sum(Payment.amount), 0.0)).scalar() or 0.0
    net_profit = total_income - total_expenses - total_payments

    pending_quotes = Quote.query.filter(
        Quote.user_id == current_user.id,
        Quote.status.in_(['draft', 'sent'])).all()
    pending_quote_count = len(pending_quotes)
    pending_quote_value = sum((q.total or 0) for q in pending_quotes)
    recent_quotes = (Quote.query.filter_by(user_id=current_user.id)
                     .order_by(Quote.created_at.desc()).limit(5).all())

    today = date.today()
    horizon = today + timedelta(days=30)
    upcoming_reminders = (Reminder.query
                          .filter(Reminder.user_id == current_user.id,
                                  Reminder.is_done.is_(False),
                                  Reminder.due_date.isnot(None),
                                  Reminder.due_date <= horizon)
                          .order_by(Reminder.due_date.asc())
                          .limit(6).all())

    return render_template(
        "home.html",
        worker_count=worker_count,
        project_count=project_count,
        active_projects=active_projects,
        active_workers=active_workers,
        customer_count=customer_count,
        total_income=total_income,
        total_expenses=total_expenses,
        net_profit=net_profit,
        pending_quote_count=pending_quote_count,
        pending_quote_value=pending_quote_value,
        recent_quotes=recent_quotes,
        upcoming_reminders=upcoming_reminders,
        today=today,
        timedelta=timedelta,
    )

@app.route('/workers', methods=['GET', 'POST'])
@login_required
def workers():
    form = WorkerForm()
    if form.validate_on_submit():
        new_worker = Worker(
            name=form.  name.data,
            contact=form.contact.data,
            daily_rate=form.daily_rate.data,
            active=form.active.data,
            phone=Worker.normalize_phone(form.phone.data),
        )
        if form.pin.data:
            new_worker.set_pin(form.pin.data)
        db.session.add(new_worker)
        db.session.commit()
        flash('Worker added successfully')
        return redirect(url_for('workers'))

    all_workers = Worker.query.all()
    return render_template('workers.html', form=form, workers=all_workers)

@app.route('/edit_worker/<int:worker_id>', methods=['GET', 'POST'])
@login_required
def edit_worker(worker_id):
    worker = Worker.query.get_or_404(worker_id)
    form = WorkerForm(obj=worker)

    if form.validate_on_submit():
        worker.name = form.name.data
        worker.contact = form.contact.data
        worker.daily_rate = form.daily_rate.data
        worker.active = form.active.data
        worker.phone = Worker.normalize_phone(form.phone.data)
        if form.pin.data:
            worker.set_pin(form.pin.data)
        db.session.commit()
        flash('Worker updated successfully.')
        return redirect(url_for('workers'))

    return render_template('edit_worker.html', form=form)

@app.route('/toggle_worker/<int:worker_id>')
@login_required
def toggle_worker(worker_id):
    worker = Worker.query.get_or_404(worker_id)
    worker.active = not worker.active
    db.session.commit()
    flash(f'Worker {"activated" if worker.active else "deactivated"} successfully.')
    return redirect(url_for('workers'))

@app.route('/delete_worker/<int:worker_id>', methods=['POST'])
@login_required
def delete_worker(worker_id):
    worker = Worker.query.get_or_404(worker_id)
    db.session.delete(worker)
    db.session.commit()
    flash('Worker deleted successfully.')
    return redirect(url_for('workers'))

# Route for Projects
def _job_customer_choices():
    custs = (Customer.query.filter_by(user_id=current_user.id)
             .order_by(Customer.name.asc()).all())
    return [(0, '— None —')] + [(c.id, c.name) for c in custs]


def _sync_project_customer_location(project, customer_id, address):
    """Set the job's customer, and find-or-create the Location from its address
    (reused when the same customer already has that address — repeat work)."""
    project.customer_id = customer_id or None
    loc = find_or_create_location(current_user.id, customer_id, address)
    project.location_id = loc.id if loc else None


@app.route('/projects', methods=['GET', 'POST'])
@login_required
def projects():
    project_form = ProjectForm()
    project_form.customer_id.choices = _job_customer_choices()
    all_projects = Project.query.all()

    if project_form.validate_on_submit() and 'add_project' in request.form:
        new_project = Project(
            name=project_form.name.data,
            description=project_form.description.data,
            address=project_form.address.data,
            start_date=project_form.start_date.data,
            status=project_form.status.data
        )
        db.session.add(new_project)
        _sync_project_customer_location(new_project, project_form.customer_id.data,
                                        project_form.address.data)
        db.session.commit()
        flash('Job added successfully')
        return redirect(url_for('project_detail', project_id=new_project.id))

    search = (request.args.get('q') or '').strip().lower()
    status_filter = (request.args.get('status') or 'all').strip()
    today = datetime.utcnow().date()
    job_rows = []

    for project in all_projects:
        contracts = list(project.contracts)
        invoices = list(project.invoices)
        contract_total = sum(c.contract_total or 0 for c in contracts)
        billed_total = sum(c.billed_to_date for c in contracts)
        paid_total = sum(
            invoice.total or 0 for invoice in invoices
            if invoice.status == 'paid')
        income_total = sum(item.amount or 0 for item in project.incomes)
        expense_total = sum(item.amount or 0 for item in project.expenses)
        labor_total = sum(item.amount or 0 for item in project.payments)
        profit = income_total - expense_total - labor_total

        attention = []
        pending_change_orders = sum(
            1 for contract in contracts for change_order in contract.change_orders
            if change_order.status == 'sent')
        if pending_change_orders:
            attention.append(
                f'{pending_change_orders} change order'
                f'{"s" if pending_change_orders != 1 else ""} awaiting approval')

        failed_inspections = sum(
            1 for permit in project.permits for inspection in permit.inspections
            if inspection.status == 'Failed')
        if failed_inspections:
            attention.append(
                f'{failed_inspections} failed inspection'
                f'{"s" if failed_inspections != 1 else ""}')

        expired_permits = sum(
            1 for permit in project.permits
            if (permit.status == 'Expired' or
                (permit.expiration_date and permit.expiration_date < today and
                 permit.status not in ('Finaled', 'Closed'))))
        if expired_permits:
            attention.append(
                f'{expired_permits} expired permit'
                f'{"s" if expired_permits != 1 else ""}')

        unscheduled = sum(
            max((contract.contract_total or 0) - contract.scheduled_total, 0)
            for contract in contracts)
        if unscheduled > 0.01:
            attention.append(f'${unscheduled:,.0f} not scheduled for billing')

        customer_name = project.customer.name if project.customer else ''
        searchable = ' '.join([
            project.name or '', customer_name, project.address or '',
        ]).lower()
        if search and search not in searchable:
            continue
        if status_filter != 'all' and project.status != status_filter:
            continue

        job_rows.append({
            'project': project,
            'customer_name': customer_name,
            'contract_total': contract_total,
            'billed_total': billed_total,
            'paid_total': paid_total,
            'profit': profit,
            'attention': attention,
            'file_count': len(project.files),
        })

    status_order = {'Active': 0, 'On Hold': 1, 'Completed': 2}
    job_rows.sort(key=lambda row: (
        status_order.get(row['project'].status, 3),
        -(row['project'].id or 0)))

    dashboard = {
        'active': sum(1 for project in all_projects if project.status == 'Active'),
        'on_hold': sum(1 for project in all_projects if project.status == 'On Hold'),
        'completed': sum(1 for project in all_projects if project.status == 'Completed'),
        'contract_total': sum(
            c.contract_total or 0 for project in all_projects
            for c in project.contracts),
        'billed_total': sum(
            c.billed_to_date for project in all_projects
            for c in project.contracts),
        'paid_total': sum(
            invoice.total or 0 for project in all_projects
            for invoice in project.invoices if invoice.status == 'paid'),
    }

    return render_template(
        'projects.html', projects=job_rows, dashboard=dashboard,
        search=request.args.get('q', ''), status_filter=status_filter)


@app.route('/projects/<int:project_id>')
@login_required
def project_detail(project_id):
    """Job hub — everything linked to one job in a single view."""
    project = Project.query.get_or_404(project_id)

    # Money (matches the dashboard: income − expenses − labor pay)
    expenses_total = sum(e.amount or 0 for e in project.expenses)
    income_total = sum(i.amount or 0 for i in project.incomes)
    labor_total = sum(p.amount or 0 for p in project.payments)
    profit = income_total - expenses_total - labor_total

    contracts = sorted(project.contracts, key=lambda c: c.id, reverse=True)
    invoices = sorted(project.invoices, key=lambda i: i.id, reverse=True)
    proposals = sorted(project.quotes, key=lambda q: q.id, reverse=True)
    change_orders = [co for c in contracts for co in c.change_orders]
    permits = sorted(project.permits, key=lambda p: p.id, reverse=True)
    documents = sorted(project.gc_documents, key=lambda d: d.id, reverse=True)

    contract_total = sum(c.contract_total or 0 for c in contracts)
    billed_total = sum(c.billed_to_date for c in contracts)
    paid_total = sum(inv.total or 0 for inv in invoices if inv.status == 'paid')

    tasks = sorted(project.tasks, key=lambda t: (t.status == 'done', t.id))
    workers = Worker.query.filter_by(active=True).order_by(Worker.name.asc()).all()

    return render_template(
        'project_detail.html', project=project,
        expenses_total=expenses_total, income_total=income_total,
        labor_total=labor_total, profit=profit,
        contracts=contracts, invoices=invoices, proposals=proposals,
        change_orders=change_orders, permits=permits, documents=documents,
        work_logs=list(project.work_logs), files=list(project.files),
        contract_total=contract_total, billed_total=billed_total, paid_total=paid_total,
        tasks=tasks, workers=workers,
        file_form=FileUploadForm(),
    )


@app.route('/projects/<int:project_id>/tasks/add', methods=['POST'])
@login_required
def add_task(project_id):
    from models import Task
    project = Project.query.get_or_404(project_id)
    description = (request.form.get('description') or '').strip()
    if not description:
        flash('Task needs a description.', 'error')
        return redirect(url_for('project_detail', project_id=project.id))
    worker_id = request.form.get('worker_id', type=int)
    due_raw = (request.form.get('due_date') or '').strip()
    due_date = None
    if due_raw:
        try:
            due_date = datetime.strptime(due_raw, '%Y-%m-%d').date()
        except ValueError:
            pass
    db.session.add(Task(project_id=project.id, worker_id=worker_id or None,
                        description=description, due_date=due_date))
    db.session.commit()
    flash('Task assigned.')
    return redirect(url_for('project_detail', project_id=project.id))


@app.route('/tasks/<int:task_id>/delete', methods=['POST'])
@login_required
def delete_task(task_id):
    from models import Task
    task = Task.query.get_or_404(task_id)
    project_id = task.project_id
    db.session.delete(task)
    db.session.commit()
    flash('Task removed.')
    return redirect(url_for('project_detail', project_id=project_id))


@app.route('/timeclock')
@login_required
def timeclock():
    from models import TimePunch
    open_punches = (TimePunch.query.filter_by(clock_out=None)
                    .order_by(TimePunch.clock_in.desc()).all())
    recent = (TimePunch.query.filter(TimePunch.clock_out.isnot(None))
              .order_by(TimePunch.clock_in.desc()).limit(50).all())
    return render_template('timeclock.html', open_punches=open_punches, recent=recent)


@app.route('/timeclock/<int:punch_id>/approve', methods=['POST'])
@login_required
def approve_punch(punch_id):
    from models import TimePunch
    p = TimePunch.query.get_or_404(punch_id)
    p.approved = not p.approved
    db.session.commit()
    flash('Punch approved.' if p.approved else 'Approval removed.')
    return redirect(url_for('timeclock'))


@app.route('/timeclock/<int:punch_id>/edit', methods=['POST'])
@login_required
def edit_punch(punch_id):
    from models import TimePunch
    p = TimePunch.query.get_or_404(punch_id)
    ci = (request.form.get('clock_in') or '').strip()
    co = (request.form.get('clock_out') or '').strip()
    if ci:
        try:
            p.clock_in = datetime.strptime(ci, '%Y-%m-%dT%H:%M')
        except ValueError:
            pass
    p.clock_out = None
    if co:
        try:
            p.clock_out = datetime.strptime(co, '%Y-%m-%dT%H:%M')
        except ValueError:
            pass
    db.session.commit()
    flash('Punch corrected.')
    return redirect(url_for('timeclock'))


@app.route('/timeclock/<int:punch_id>/delete', methods=['POST'])
@login_required
def delete_punch(punch_id):
    from models import TimePunch
    p = TimePunch.query.get_or_404(punch_id)
    db.session.delete(p)
    db.session.commit()
    flash('Punch deleted.')
    return redirect(url_for('timeclock'))


def _archive_job_documents(project):
    """Save the job's current proposal/contract/invoice/waiver PDFs into its
    Files (static/uploads) so all paperwork lives with the job. Skips ones
    already filed (by filename). Returns how many were newly saved."""
    import os
    from quote_routes import build_quote_pdf
    from invoice_routes import build_invoice_pdf
    from contract_routes import build_contract_pdf
    from waiver_routes import build_gc_document_pdf

    folder = app.config['UPLOAD_FOLDER']
    existing = {f.filename for f in project.files}
    saved = 0

    def _save(buf, filename, category):
        nonlocal saved
        if filename in existing:
            return
        path = os.path.join(folder, filename)
        with open(path, 'wb') as fh:
            fh.write(buf.read())
        db.session.add(ProjectFile(project_id=project.id, filename=filename,
                                   filepath=path, category=category, note='Auto-filed'))
        existing.add(filename)
        saved += 1

    for q in project.quotes:
        _save(build_quote_pdf(q), f'proposal-{q.id}.pdf', 'proposal')
    for c in project.contracts:
        _save(build_contract_pdf(c), f'contract-{c.number or c.id}.pdf', 'contract')
    for inv in project.invoices:
        _save(build_invoice_pdf(inv), f'invoice-{inv.number or inv.id}.pdf', 'invoice')
    for d in project.gc_documents:
        _save(build_gc_document_pdf(d), f'{d.doc_kind}-{d.id}.pdf', 'misc')
    return saved


@app.route('/projects/<int:project_id>/file-documents', methods=['POST'])
@login_required
def project_file_documents(project_id):
    project = Project.query.get_or_404(project_id)
    filed = _archive_job_documents(project)
    db.session.commit()
    flash(f'{filed} document(s) filed to this job.' if filed
          else 'All documents are already filed.')
    return redirect(url_for('project_detail', project_id=project.id))


@app.route('/projects/<int:project_id>/closeout', methods=['POST'])
@login_required
def project_closeout(project_id):
    from datetime import date
    from models import GcDocument
    project = Project.query.get_or_404(project_id)
    owner = project.customer.name if project.customer else None
    paid_total = sum(i.total or 0 for i in project.invoices if i.status == 'paid')

    db.session.add(GcDocument(
        user_id=current_user.id, doc_kind='completion', project_id=project.id,
        customer_id=project.customer_id, owner_name=owner,
        property_address=project.address, through_date=date.today(),
        notes=project.description or project.name))
    db.session.add(GcDocument(
        user_id=current_user.id, doc_kind='unconditional_final', project_id=project.id,
        customer_id=project.customer_id, owner_name=owner,
        property_address=project.address, amount=paid_total, through_date=date.today()))

    for contract in project.contracts:
        contract.status = 'completed'
    project.status = 'Completed'
    db.session.flush()
    filed = _archive_job_documents(project)
    db.session.commit()
    flash('Job closed out — completion certificate and final lien waiver generated, '
          f'{filed} document(s) filed, contract and job marked complete.')
    return redirect(url_for('project_detail', project_id=project.id))


@app.route('/upload_file/<int:project_id>', methods=['POST'])
@login_required
def upload_file(project_id):
    form= FileUploadForm()
    if form.validate_on_submit():
        file = form.file.data
        filename = save_uploaded_file(file)
        if not filename:
            flash('File type not allowed.', 'error')
            return redirect(url_for('project_detail', project_id=project_id))

        new_file = ProjectFile(
            project_id=project_id,
            filename=filename,
            filepath=os.path.join(app.config['UPLOAD_FOLDER'], filename),
            category=form.category.data,
            note=form.note.data if form.category.data == 'misc' else None
        )
        db.session.add(new_file)
        db.session.commit()
        flash('File uploaded successfully')
    return redirect(url_for('project_detail', project_id=project_id))


@app.route('/edit_project/<int:project_id>', methods=['GET', 'POST'])
@login_required
def edit_project(project_id):
    project = Project.query.get_or_404(project_id)
    form = ProjectForm(obj=project)
    form.customer_id.choices = _job_customer_choices()
    if request.method == 'GET':
        form.customer_id.data = project.customer_id or 0

    if form.validate_on_submit():
        project.name = form.name.data
        project.description = form.description.data
        project.address = form.address.data
        project.start_date = form.start_date.data
        project.status = form.status.data
        _sync_project_customer_location(project, form.customer_id.data, form.address.data)
        db.session.commit()
        flash('Job updated successfully.')
        return redirect(url_for('project_detail', project_id=project.id))

    return render_template('edit_project.html', form=form, project=project)

@app.route('/delete_project/<int:project_id>', methods=['POST'])
@login_required
@owner_required
def delete_project(project_id):
    project = Project.query.get_or_404(project_id)
    for f in project.files:
        if f.filepath and os.path.exists(f.filepath):
            os.remove(f.filepath)
    db.session.delete(project)
    db.session.commit()
    flash('Project deleted successfully.')
    return redirect(url_for('projects'))

@app.route('/delete_file/<int:file_id>', methods=['POST'])
@login_required
def delete_file(file_id):
    file = ProjectFile.query.get_or_404(file_id)
    if file.filepath and os.path.exists(file.filepath):
        os.remove(file.filepath)
    db.session.delete(file)
    db.session.commit()
    flash("File deleted.")
    return redirect(url_for('projects'))


@app.route('/edit_file/<int:file_id>', methods=['GET', 'POST'])
@login_required
def edit_file(file_id):
    file = ProjectFile.query.get_or_404(file_id)
    form = FileUploadForm(obj=file)

    if form.validate_on_submit():
        file.category = form.category.data
        file.note = form.note.data
        db.session.commit()
        flash("File updated.")
        return redirect(url_for('projects'))

    return render_template('edit_file.html', form=form, file=file)

# Route for Work Logs
@app.route('/work_logs', methods=['GET', 'POST'])
@login_required
def work_logs():
    form = WorkLogForm()
    filter_form = WorkLogFilterForm()

    # Set dropdown choices
    workers = Worker.query.all()
    projects = Project.query.all()
    form.worker_id.choices = [(w.id, w.name) for w in workers]
    form.project_id.choices = [(p.id, p.name) for p in projects]
    filter_form.worker_id.choices = [(-1, 'All')] + [(w.id, w.name) for w in workers]
    filter_form.project_id.choices = [(-1, 'All')] + [(p.id, p.name) for p in projects]


    # Handle new log submission
    if form.validate_on_submit() and 'submit' in request.form:
        new_log = WorkLog(
            worker_id=form.worker_id.data,
            project_id=form.project_id.data,
            start_date=form.start_date.data,
            end_date=form.end_date.data,
            days_worked=form.days_worked.data,
            note=form.note.data
        )
        db.session.add(new_log)
        db.session.flush()

        pay_msg = ''
        if form.create_payment.data:
            worker = Worker.query.get(form.worker_id.data)
            amount = (form.days_worked.data or 0) * ((worker.daily_rate or 0) if worker else 0)
            if amount > 0:
                db.session.add(Payment(
                    worker_id=form.worker_id.data,
                    project_id=form.project_id.data,
                    amount=amount,
                    payment_date=form.end_date.data or form.start_date.data,
                    method='Labor',
                    note='Auto-recorded from work log',
                    work_log_id=new_log.id,
                ))
                pay_msg = f' · ${amount:,.2f} labor payment recorded'

        db.session.commit()
        flash('Work log added' + pay_msg)
        return redirect(url_for('work_logs'))

    # Handle filtering
    logs_query = WorkLog.query
    if filter_form.validate_on_submit() and 'filter' in request.form:
        if filter_form.worker_id.data != -1:
            logs_query = logs_query.filter_by(worker_id=filter_form.worker_id.data)
        if filter_form.project_id.data != -1:
            logs_query = logs_query.filter_by(project_id=filter_form.project_id.data)
        if filter_form.start_date.data:
            logs_query = logs_query.filter(WorkLog.start_date >= filter_form.start_date.data)
        if filter_form.end_date.data:
            logs_query = logs_query.filter(WorkLog.end_date <= filter_form.end_date.data)

    logs = logs_query.order_by(WorkLog.start_date.desc()).all()

    return render_template(
        'work_logs.html',
        form=form,
        filter_form=filter_form,
        logs=logs
    )




@app.route('/work_logs/edit/<int:log_id>', methods=['GET', 'POST'])
@login_required
def edit_work_log(log_id):
    log = WorkLog.query.get_or_404(log_id)
    form = WorkLogForm(obj=log)

    form.worker_id.choices = [(w.id, w.name) for w in Worker.query.all()]
    form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]

    if form.validate_on_submit():
        log.worker_id = form.worker_id.data
        log.project_id = form.project_id.data
        log.start_date = form.start_date.data
        log.end_date = form.end_date.data
        log.days_worked = form.days_worked.data
        log.note = form.note.data
        db.session.commit()
        flash('Work log updated successfully')
        return redirect(url_for('work_logs'))

    return render_template('edit_work_log.html', form=form, log=log)



@app.route('/work_logs/delete/<int:log_id>', methods=['POST'])
@login_required
def delete_work_log(log_id):
    log = WorkLog.query.get_or_404(log_id)
    db.session.delete(log)
    db.session.commit()
    flash('Work log deleted successfully.')
    return redirect(url_for('work_logs'))


@app.route('/payments', methods=['GET', 'POST'])
@login_required
def payments():
    form = PaymentForm()
    filter_form = PaymentFilterForm()
    form.worker_id.choices = [(0, '---')] + [(w.id, w.name) for w in Worker.query.all()]
    form.project_id.choices = [(0, '---')] + [(p.id, p.name) for p in Project.query.all()]
    
    # Handle submission
    if form.validate_on_submit():
        # Handle receipt upload
        filename = None
        filepath = None
        if form.receipt.data:
            receipt_file = form.receipt.data
            filename = save_uploaded_file(receipt_file)
            if not filename:
                flash('File type not allowed.', 'error')
                return redirect(url_for('payments'))
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        new_payment = Payment(
            worker_id=form.worker_id.data if form.worker_id.data != 0 else None,
            project_id=form.project_id.data if form.project_id.data != 0 else None,
            amount=form.amount.data,
            payment_date=form.payment_date.data,
            method=form.method.data,
            note=form.note.data,
            receipt_filename=filename,
            receipt_filepath=filepath
        )
        db.session.add(new_payment)
        db.session.commit()
        flash('Payment recorded successfully')
        return redirect(url_for('payments'))

    # Handle filter
    payments_query = Payment.query

    worker_id = request.args.get('worker_id', type=int)
    if worker_id:
        payments_query = payments_query.filter_by(worker_id=worker_id)

    start_date = request.args.get('start_date')
    if start_date:
        payments_query = payments_query.filter(Payment.payment_date >= start_date)

    end_date = request.args.get('end_date')
    if end_date:
        payments_query = payments_query.filter(Payment.payment_date <= end_date)

    min_amount = request.args.get('min_amount', type=float)
    if min_amount is not None:
        payments_query = payments_query.filter(Payment.amount >= min_amount)

    max_amount = request.args.get('max_amount', type=float)
    if max_amount is not None:
        payments_query = payments_query.filter(Payment.amount <= max_amount)

    filtered_payments = payments_query.order_by(Payment.payment_date.desc()).all()

    return render_template('payments.html',
        form=form,
        payments=filtered_payments,
        filter_form=filter_form,
        workers=Worker.query.all()
    )

@app.route('/payments/edit/<int:payment_id>', methods=['GET', 'POST'])
@login_required
def edit_payment(payment_id):
    payment = Payment.query.get_or_404(payment_id)
    form = PaymentForm(obj=payment)
    form.worker_id.choices = [(0, '---')] + [(w.id, w.name) for w in Worker.query.all()]
    form.project_id.choices = [(0, '---')] + [(p.id, p.name) for p in Project.query.all()]

    if form.validate_on_submit():
        payment.worker_id = form.worker_id.data if form.worker_id.data != 0 else None
        payment.project_id = form.project_id.data if form.project_id.data != 0 else None
        payment.amount = form.amount.data
        payment.payment_date = form.payment_date.data
        payment.method = form.method.data
        payment.note = form.note.data
        
        # Handle new receipt upload
        if form.receipt.data:
            # Delete old receipt if exists
            if payment.receipt_filename:
                try:
                    old_filepath = os.path.join(app.config['UPLOAD_FOLDER'], payment.receipt_filename)
                    if os.path.exists(old_filepath):
                        os.remove(old_filepath)
                except FileNotFoundError:
                    pass
            
            receipt_file = form.receipt.data
            filename = save_uploaded_file(receipt_file)
            if not filename:
                flash('File type not allowed.', 'error')
                return redirect(url_for('payments'))
            payment.receipt_filename = filename
            payment.receipt_filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        db.session.commit()
        flash('Payment updated successfully.')
        return redirect(url_for('payments'))

    return render_template('edit_payment.html', form=form, payment=payment)
    

@app.route('/payments/delete/<int:payment_id>', methods=['POST'])
@login_required
@owner_required
def delete_payment(payment_id):
    payment = Payment.query.get_or_404(payment_id)
    
    # Delete receipt file if exists
    if payment.receipt_filename:
        try:
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], payment.receipt_filename)
            if os.path.exists(filepath):
                os.remove(filepath)
        except FileNotFoundError:
            pass
    
    db.session.delete(payment)
    db.session.commit()
    flash('Payment deleted successfully.')
    return redirect(url_for('payments'))


@app.route('/expenses', methods=['GET', 'POST'])
@login_required
def expenses():
    form = ExpenseForm()
    form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]

    if form.validate_on_submit():
        filename = None
        filepath = None
        if form.receipt.data:
            receipt_file = form.receipt.data
            filename = save_uploaded_file(receipt_file)
            if not filename:
                flash('File type not allowed.', 'error')
                return redirect(url_for('expenses'))
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        new_expense = Expense(
            project_id=form.project_id.data,
            description=form.description.data,
            amount=form.amount.data,
            category=form.category.data,
            date=form.date.data,
            note=form.note.data,
            receipt_filename=filename
        )
        db.session.add(new_expense)
        db.session.commit()
        flash('Expense added successfully')
        return redirect(url_for('expenses'))

    all_expenses = Expense.query.order_by(Expense.date.desc()).all()
    return render_template('expenses.html', form=form, expenses=all_expenses)

@app.route('/expenses/delete/<int:expense_id>', methods=['POST'])
@login_required
@owner_required
def delete_expense(expense_id):
    expense = Expense.query.get_or_404(expense_id)
    if expense.receipt_filename:
        try:
            os.remove(os.path.join(app.config['UPLOAD_FOLDER'], expense.receipt_filename))
        except FileNotFoundError:
            pass  # If file doesn't exist, silently ignore
    db.session.delete(expense)
    db.session.commit()
    flash('Expense deleted successfully')
    return redirect(url_for('expenses'))



@app.route('/expenses/edit/<int:expense_id>', methods=['GET', 'POST'])
@login_required
def edit_expense(expense_id):
    expense = Expense.query.get_or_404(expense_id)
    form = ExpenseForm(obj=expense)
    form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]

    if form.validate_on_submit():
        # Update fields
        expense.project_id = form.project_id.data
        expense.description = form.description.data
        expense.amount = form.amount.data
        expense.category = form.category.data
        expense.date = form.date.data
        expense.note = form.note.data

        # Handle new receipt upload
        if form.receipt.data:
            if expense.receipt_filename:
                try:
                    os.remove(os.path.join(app.config['UPLOAD_FOLDER'], expense.receipt_filename))
                except FileNotFoundError:
                    pass
            receipt_file = form.receipt.data
            filename = save_uploaded_file(receipt_file)
            if not filename:
                flash('File type not allowed.', 'error')
                return redirect(url_for('expenses'))
            expense.receipt_filename = filename
            expense.receipt_filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        db.session.commit()
        flash('Expense updated successfully')
        return redirect(url_for('expenses'))

    return render_template('edit_expense.html', form=form)

@app.route('/profitability')
@login_required
def profitability():
    projects = Project.query.all()
    data = []

    for project in projects:
        total_expenses = sum(exp.amount for exp in project.expenses)
        total_income = sum(inc.amount for inc in project.incomes)  # 👈 use incomes from relationship
        profit = total_income - total_expenses

        data.append({
            'project': project.name,
            'expenses': total_expenses,
            'income': total_income,
            'profit': profit
        })

    return render_template('profitability.html', data=data)


@app.route('/income', methods=['GET', 'POST'])
@login_required
def income():
    form = IncomeForm()
    form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]

    if form.validate_on_submit():
        new_income = Income(
            project_id=form.project_id.data,
            amount=form.amount.data,
            source=form.source.data,
            date=form.date.data,
            note=form.note.data
        )
        db.session.add(new_income)
        db.session.commit()
        flash('Income added successfully')
        return redirect(url_for('income'))

    all_income = Income.query.order_by(Income.date.desc()).all()
    return render_template('income.html', form=form, incomes=all_income)


@app.route('/income/edit/<int:income_id>', methods=['GET', 'POST'])
@login_required
def edit_income(income_id):
    income = Income.query.get_or_404(income_id)
    form = IncomeForm(obj=income)
    form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]

    if form.validate_on_submit():
        income.project_id = form.project_id.data
        income.amount = form.amount.data
        income.source = form.source.data
        income.date = form.date.data
        income.note = form.note.data
        db.session.commit()
        flash('Income updated successfully.')
        return redirect(url_for('income'))

    return render_template('edit_income.html', form=form, income=income)

@app.route('/income/delete/<int:income_id>', methods=['POST'])
@login_required
@owner_required
def delete_income(income_id):
    income = Income.query.get_or_404(income_id)
    db.session.delete(income)
    db.session.commit()
    flash('Income deleted successfully.')
    return redirect(url_for('income'))

# Document Generation Routes
def parse_scope_items(scope_text):
    """
    Parse scope of work text into individual items.
    Splits on periods, semicolons, or newlines.
    """
    if not scope_text:
        return []
    
    # Replace semicolons and newlines with periods for uniform splitting
    scope_text = scope_text.replace(';', '.')
    scope_text = scope_text.replace('\n', '.')
    
    # Split by periods and clean up
    items = [item.strip() for item in scope_text.split('.') if item.strip()]
    
    return items

# EXACT FIX FOR YOUR app.py
# Replace your document_form route (around line 375) with this:

@app.route('/documents/create', methods=['GET', 'POST'])
@login_required
def document_form():
    form = DocumentForm()
    
    if request.method == 'POST':
        if not form.validate_on_submit():
            app.logger.error(f"Validation failed with errors: {form.errors}")
            # Flash each error to user
            for field_name, errors in form.errors.items():
                for error in errors:
                    flash(f'{field_name}: {error}', 'error')
    
    if form.validate_on_submit():
        try:
            app.logger.info("Form validated successfully, storing session data")
            
            # Store form data in session
            session['document_data'] = {
                'doc_type': form.doc_type.data,
                'invoice_number': form.invoice_number.data,
                'date': form.date.data.strftime('%m-%d-%y'),
                'client_name': form.client_name.data,
                'client_address': form.client_address.data,
                'project_name': form.project_name.data,
                'scope_of_work': form.scope_of_work.data,
                'materials_notes': form.materials_notes.data if form.materials_notes.data else '',
                'total_price': float(form.total_price.data),
                'po_number': form.po_number.data if form.po_number.data else None
            }
            
            # CRITICAL: Mark session as modified
            session.modified = True
            
            app.logger.info(f"Session data stored for client: {form.client_name.data}")
            app.logger.info(f"Redirecting to view_document")
            
            flash('Document generated successfully!', 'success')
            return redirect(url_for('view_document'))
            
        except Exception as e:
            app.logger.error(f"Error storing session data: {e}", exc_info=True)
            flash(f'Error generating document: {str(e)}', 'error')
            return render_template('documents/form.html', form=form)
    
    # If we reach here with POST, it means validation failed
    if request.method == 'POST':
        app.logger.warning("Rendering form with validation errors")
    
    return render_template('documents/form.html', form=form)

@app.route('/documents/view')
@login_required
def view_document():
    # Get document data from session
    doc_data = session.get('document_data')

    if not doc_data:
        flash('No document data found. Please create a new document.')
        return redirect(url_for('document_form'))

    try:
        # Parse scope items from the raw scope_of_work text
        scope_items = parse_scope_items(doc_data.get('scope_of_work', ''))

        # IMPORTANT: pass each field as a keyword argument,
        # not a single dict or positional arg.
        return render_template(
            'documents/document.html',
            doc_type=doc_data['doc_type'],
            invoice_number=doc_data['invoice_number'],
            date=doc_data['date'],
            client_name=doc_data['client_name'],
            client_address=doc_data['client_address'],
            project_name=doc_data['project_name'],
            scope_items=scope_items,
            materials_notes=doc_data.get('materials_notes', ''),
            total_price=doc_data['total_price'],
            po_number=doc_data.get('po_number')
        )

    except Exception as e:
        app.logger.error(f"Error in view_document: {e}", exc_info=True)
        flash('There was a problem generating the document.', 'error')
        return redirect(url_for('document_form'))


@app.route('/documents/download')
@login_required
def download_document():
    """Generate and download the document as a Word file"""
    doc_data = session.get('document_data')
    
    if not doc_data:
        flash('No document data found. Please create a new document.')
        return redirect(url_for('document_form'))
    
    # Parse scope items
    scope_items = parse_scope_items(doc_data['scope_of_work'])
    
    # Create Word document
    doc = Document()
    
    # Set up default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add header section with company info
    header_section = doc.add_paragraph()
    header_section.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    company_name = header_section.add_run('Optimal Solutions\n')
    company_name.bold = True
    company_name.font.size = Pt(14)
    
    header_section.add_run('1204 W County Line Rd\n')
    header_section.add_run('Beecher IL, 60401\n')
    header_section.add_run('708-825-4774\n')
    header_section.add_run('MendezConstruction39@yahoo.com\n')
    
    # Add horizontal line
    doc.add_paragraph('_' * 80)
    
    # Document title and info
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(doc_data['doc_type'])
    title_run.bold = True
    title_run.font.size = Pt(18)
    
    # Invoice/Date info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_para.add_run(f"INVOICE # {doc_data['invoice_number']}\n").bold = True
    info_para.add_run(f"DATE {doc_data['date']}\n").bold = True
    if doc_data.get('po_number'):
        info_para.add_run(f"P.O. # {doc_data['po_number']}\n").bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Client information
    to_para = doc.add_paragraph()
    to_para.add_run('TO\n').bold = True
    to_para.add_run(f"{doc_data['client_name']}\n").bold = True
    to_para.add_run(f"{doc_data['client_address']}\n")
    
    doc.add_paragraph()  # Spacing
    
    # Project name
    for_para = doc.add_paragraph()
    for_para.add_run('FOR\n').bold = True
    for_para.add_run(f"{doc_data['project_name']}\n")
    
    doc.add_paragraph()  # Spacing
    
    # Scope of work table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Description'
    header_cells[1].text = 'Amount'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add scope items
    for item in scope_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = ''
    
    # Add spacing row
    table.add_row()
    
    # Materials notes if provided
    if doc_data.get('materials_notes'):
        materials_row = table.add_row().cells
        materials_para = materials_row[0].paragraphs[0]
        materials_para.add_run('Materials:\n').bold = True
        materials_para.add_run(doc_data['materials_notes'])
        materials_row[0].merge(materials_row[1])
        table.add_row()  # Spacing
    
    # Total row
    total_row = table.add_row().cells
    total_row[0].text = 'Total'
    total_row[1].text = f"${doc_data['total_price']:.2f}"
    
    # Make total row bold
    for cell in total_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.add_run('Make all checks payable to Optimal Solutions\n').bold = True
    footer_para.add_run('Payment is due once work is completed\n')
    footer_para.add_run('\nIf you have any questions concerning this ')
    footer_para.add_run(doc_data['doc_type'].lower())
    footer_para.add_run(', contact:\n')
    footer_para.add_run('Mario Mendez | 708-825-4774 | MendezConstruction39@yahoo.com\n')
    
    doc.add_paragraph()  # Spacing
    doc.add_paragraph('_' * 80)
    
    # Thank you
    thanks_para = doc.add_paragraph()
    thanks_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thanks_run = thanks_para.add_run('THANK YOU FOR YOUR BUSINESS!')
    thanks_run.bold = True
    thanks_run.font.size = Pt(14)
    
    # Save to BytesIO object
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Generate filename
    filename = f"{doc_data['doc_type']}_{doc_data['invoice_number']}_{doc_data['client_name'].replace(' ', '_')}.docx"
    
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )
#Monitoring route
logging.basicConfig(level=logging.INFO)
db_logger = logging.getLogger('database')

class DatabaseHealthMonitor:
    def __init__(self, db):
        self.db = db
        self.last_health_check = None
        self.consecutive_failures = 0
        self.max_failures = 3
        
    def check_connection_health(self):
        """Proactive health check to catch issues early"""
        try:
            with self.db.engine.connect() as connection:
                connection.execute(text("SELECT 1"))
                self.consecutive_failures = 0
                self.last_health_check = datetime.utcnow()
                return True
        except (OperationalError, DisconnectionError) as e:
            self.consecutive_failures += 1
            db_logger.warning(f"Database health check failed: {e}")
            return False
    
    def should_check_health(self):
        """Check if we need to run a health check"""
        if not self.last_health_check:
            return True
        return datetime.utcnow() - self.last_health_check > timedelta(minutes=5)

# Initialize the monitor
health_monitor = DatabaseHealthMonitor(db)

def database_retry(max_retries=3, delay=1, backoff=2):
    """Decorator to retry database operations with exponential backoff"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            # Check connection health periodically
            if health_monitor.should_check_health():
                health_monitor.check_connection_health()
            
            for attempt in range(max_retries):
                try:
                    return func(*args, **kwargs)
                except (OperationalError, DisconnectionError) as e:
                    error_msg = str(e)
                    
                    # Log the specific error
                    db_logger.error(f"Database error on attempt {attempt + 1}: {error_msg}")
                    
                    # Check if it's a recoverable SSL/connection error
                    recoverable_errors = [
                        "SSL connection has been closed unexpectedly",
                        "server closed the connection unexpectedly",
                        "connection not open",
                        "connection already closed",
                        "could not connect to server"
                    ]
                    
                    is_recoverable = any(err in error_msg.lower() for err in recoverable_errors)
                    
                    if is_recoverable and attempt < max_retries - 1:
                        # Rollback any pending transaction
                        try:
                            db.session.rollback()
                        except:
                            pass
                        
                        # Wait before retrying (exponential backoff)
                        wait_time = delay * (backoff ** attempt)
                        db_logger.info(f"Retrying in {wait_time} seconds...")
                        time.sleep(wait_time)
                        continue
                    
                    # If not recoverable or max retries reached, re-raise
                    raise e
                except Exception as e:
                    # Non-database errors should not be retried
                    raise e
            
            return func(*args, **kwargs)
        return wrapper
    return decorator

# Add this route to your app for health monitoring
@app.route('/health')
def health_check():
    """Health endpoint for monitoring systems"""
    try:
        # Test database connection
        with db.engine.connect() as connection:
            connection.execute(text("SELECT 1"))
        
        return {
            'status': 'healthy',
            'timestamp': datetime.utcnow().isoformat(),
            'database': 'connected'
        }, 200
    except Exception as e:
        app.logger.error(f"Health check failed: {e}", exc_info=True)
        return {
            'status': 'unhealthy',
            'timestamp': datetime.utcnow().isoformat(),
            'database': 'disconnected'
        }, 503

# Add error handlers for better monitoring
@app.errorhandler(OperationalError)
def handle_database_error(error):
    error_msg = str(error)
    db_logger.error(f"Database operational error: {error_msg}")
    
    try:
        db.session.rollback()
    except:
        pass
    
    # Check if it's an SSL connection error
    if "SSL connection has been closed unexpectedly" in error_msg:
        flash('Connection temporarily lost. Please try again.')
    else:
        flash('Database error. Please try again in a moment.')
    
    return redirect(url_for('home'))

@app.errorhandler(DisconnectionError)
def handle_disconnection_error(error):
    db_logger.error(f"Database disconnection error: {error}")
    
    try:
        db.session.rollback()
    except:
        pass
    
    flash('Connection lost. Please try again.')
    return redirect(url_for('home'))

if __name__ == '__main__':
        with app.app_context():
            db.create_all()
        app.run(debug=True)
